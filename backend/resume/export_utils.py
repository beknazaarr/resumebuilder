from io import BytesIO
from django.template.loader import render_to_string
from weasyprint import HTML
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_html_from_template(resume):
    """
    Генерирует HTML с подстановкой данных резюме в шаблон
    """
    template = resume.template
    html = template.html_structure
    css = template.css_styles
    
    # Получаем данные
    personal_info = getattr(resume, 'personal_info', None)
    education = resume.education.all()
    work_experience = resume.work_experience.all()
    skills = resume.skills.all()
    achievements = resume.achievements.all()
    languages = resume.languages.all()
    
    # ============= ИСПРАВЛЕНИЕ ФОТО =============
    if resume.photo:
        # Получаем ПОЛНЫЙ путь к фото
        try:
            # Если это объект ImageField
            if hasattr(resume.photo, 'path'):
                import base64
                with open(resume.photo.path, 'rb') as img_file:
                    img_data = base64.b64encode(img_file.read()).decode('utf-8')
                    # Определяем тип изображения
                    ext = resume.photo.name.split('.')[-1].lower()
                    mime_type = f'image/{ext}' if ext in ['jpg', 'jpeg', 'png', 'webp'] else 'image/jpeg'
                    photo_url = f'data:{mime_type};base64,{img_data}'
            else:
                photo_url = str(resume.photo)
        except Exception as e:
            print(f"Error loading photo: {e}")
            photo_url = ''
        
        if photo_url:
            html = html.replace('{{photo}}', photo_url)
            html = html.replace('{{#if photo}}', '')
            html = html.replace('{{else}}', '<!--')
            html = html.replace('{{/if}}', '-->')
        else:
            # Убираем блок с фото
            html = html.replace('{{#if photo}}', '<!--')
            html = html.replace('{{else}}', '')
            html = html.replace('{{/if}}', '-->')
            html = html.replace('{{photo}}', '')
    else:
        html = html.replace('{{#if photo}}', '<!--')
        html = html.replace('{{else}}', '')
        html = html.replace('{{/if}}', '-->')
        html = html.replace('{{photo}}', '')
    
    # Подстановка личной информации
    if personal_info:
        html = html.replace('{{full_name}}', personal_info.full_name or '')
        html = html.replace('{{email}}', personal_info.email or '')
        html = html.replace('{{phone}}', personal_info.phone or '')
        html = html.replace('{{address}}', personal_info.address or '')
        html = html.replace('{{linkedin}}', personal_info.linkedin or '')
        html = html.replace('{{website}}', personal_info.website or '')
        html = html.replace('{{summary}}', personal_info.summary or '')
    else:
        html = html.replace('{{full_name}}', 'Ваше Имя')
        html = html.replace('{{email}}', 'email@example.com')
        html = html.replace('{{phone}}', '+X XXX XXX XXXX')
        html = html.replace('{{address}}', '')
        html = html.replace('{{linkedin}}', '')
        html = html.replace('{{website}}', '')
        html = html.replace('{{summary}}', '')
    
    # Подстановка опыта работы
    work_html = ''
    for exp in work_experience:
        end_date = 'настоящее время' if exp.is_current else format_date_for_export(exp.end_date)
        work_html += f'''
        <div class="experience-item">
            <div class="experience-header">
                <div class="experience-title">{exp.position}</div>
                <div class="experience-date">{format_date_for_export(exp.start_date)} - {end_date}</div>
            </div>
            <div class="experience-company">{exp.company}</div>
            {f'<div class="experience-description">{exp.description}</div>' if exp.description else ''}
        </div>
        '''
    html = html.replace('{{work_experience}}', work_html)
    
    # Подстановка образования
    edu_html = ''
    for edu in education:
        end_date = format_date_for_export(edu.end_date) if edu.end_date else 'настоящее время'
        edu_html += f'''
        <div class="education-item">
            <div class="education-header">
                <div class="education-title">{edu.degree} - {edu.field_of_study}</div>
                <div class="education-date">{format_date_for_export(edu.start_date)} - {end_date}</div>
            </div>
            <div class="education-school">{edu.institution}</div>
            {f'<div class="education-description">{edu.description}</div>' if edu.description else ''}
        </div>
        '''
    html = html.replace('{{education}}', edu_html)
    
    # Подстановка навыков
    skills_html = '<ul style="list-style: none; padding: 0; margin: 0;">'
    for skill in skills:
        level_display = dict(skill.LEVEL_CHOICES).get(skill.level, skill.level)
        skills_html += f'<li style="margin-bottom: 0.3rem;">• {skill.name} ({level_display})</li>'
    skills_html += '</ul>'
    html = html.replace('{{skills}}', skills_html)
    
    # Подстановка достижений
    achievements_html = ''
    for ach in achievements:
        date_str = f'<div class="date">{format_date_for_export(ach.date)}</div>' if ach.date else ''
        achievements_html += f'''
        <div class="reward-item">
            {date_str}
            <div class="title">{ach.title}</div>
            {f'<div class="description">{ach.description}</div>' if ach.description else ''}
        </div>
        '''
    html = html.replace('{{achievements}}', achievements_html)
    html = html.replace('{{rewards}}', achievements_html)
    
    # Подстановка языков
    languages_html = ''
    for lang in languages:
        level_display = dict(lang.PROFICIENCY_CHOICES).get(lang.proficiency_level, lang.proficiency_level)
        languages_html += f'<div style="margin-bottom: 0.3rem;">{lang.language} - {level_display}</div>'
    html = html.replace('{{languages}}', languages_html)
    
    # ============= КОМПАКТНЫЙ CSS ДЛЯ ОДНОЙ СТРАНИЦЫ =============
    pdf_css = f'''
    {css}
    
    /* PDF настройки */
    @page {{
        size: A4;
        margin: 0.5cm; /* Минимальные отступы */
    }}
    
    body {{
        margin: 0;
        padding: 0;
        width: 100%;
        max-width: 210mm;
        font-size: 9pt; /* Уменьшен шрифт */
        line-height: 1.2; /* Компактная высота строк */
    }}
    
    * {{
        box-sizing: border-box;
        page-break-inside: avoid;
    }}
    
    /* Компактные заголовки */
    h1 {{ 
        font-size: 16pt; 
        margin: 0 0 0.2em 0;
        line-height: 1.1;
    }}
    
    h2 {{ 
        font-size: 12pt; 
        margin: 0.3em 0 0.2em 0;
        line-height: 1.1;
    }}
    
    h3 {{ 
        font-size: 10pt; 
        margin: 0.2em 0;
        line-height: 1.1;
    }}
    
    /* Фото компактное */
    img {{
        max-width: 80px !important; /* Уменьшили с 150px */
        max-height: 80px !important;
        border-radius: 50%;
        object-fit: cover;
        page-break-inside: avoid;
    }}
    
    /* Компактные секции */
    .experience-item,
    .education-item,
    .reward-item {{
        page-break-inside: avoid;
        margin-bottom: 0.3em; /* Минимальный отступ */
        padding: 0.2em 0;
    }}
    
    .experience-header,
    .education-header {{
        margin-bottom: 0.1em;
    }}
    
    .experience-description,
    .education-description {{
        margin-top: 0.1em;
        font-size: 8pt; /* Еще меньше */
        line-height: 1.2;
    }}
    
    /* Убираем лишние отступы */
    ul, ol {{
        margin: 0.2em 0;
        padding-left: 1.2em;
    }}
    
    li {{
        margin-bottom: 0.1em;
    }}
    
    p {{
        margin: 0.2em 0;
    }}
    
    /* Компактный контейнер */
    .pdf-container {{
        width: 100%;
        max-width: 100%;
        padding: 0;
        margin: 0;
    }}
    
    /* Убираем разрывы страниц */
    div, section, article {{
        page-break-inside: avoid;
    }}
    '''
    
    # Оборачиваем в полный HTML
    full_html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            {pdf_css}
        </style>
    </head>
    <body>
        <div class="pdf-container">
            {html}
        </div>
    </body>
    </html>
    '''
    
    return full_html


def format_date_for_export(date):
    """Форматирование даты для экспорта"""
    if not date:
        return ''
    try:
        from datetime import datetime
        if isinstance(date, str):
            date = datetime.strptime(date, '%Y-%m-%d')
        return date.strftime('%B %Y')
    except:
        return str(date)


def generate_pdf(resume):
    """Генерация PDF из резюме с учётом шаблона"""
    
    # Проверяем, есть ли у резюме шаблон
    if resume.template and resume.template.html_structure:
        # Используем HTML из шаблона
        html_content = generate_html_from_template(resume)
    else:
        # Используем дефолтный шаблон
        html_content = render_to_string('resume/pdf_template.html', {
            'resume': resume,
            'personal_info': getattr(resume, 'personal_info', None),
            'education': resume.education.all(),
            'work_experience': resume.work_experience.all(),
            'skills': resume.skills.all(),
            'achievements': resume.achievements.all(),
            'languages': resume.languages.all(),
        })
    
    # Генерируем PDF
    pdf_file = BytesIO()
    HTML(string=html_content).write_pdf(pdf_file)
    pdf_file.seek(0)
    
    return pdf_file


def generate_docx(resume):

    """Генерация DOCX из резюме"""

    doc = Document()
    
    # Настройки документа
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Личная информация
    if hasattr(resume, 'personal_info'):
        personal_info = resume.personal_info
        
        # Имя
        heading = doc.add_heading(personal_info.full_name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Контакты
        contact_info = []
        if personal_info.phone:
            contact_info.append(f"Телефон: {personal_info.phone}")
        if personal_info.email:
            contact_info.append(f"Email: {personal_info.email}")
        if personal_info.address:
            contact_info.append(f"Адрес: {personal_info.address}")
        
        if contact_info:
            p = doc.add_paragraph(" | ".join(contact_info))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ссылки
        links = []
        if personal_info.linkedin:
            links.append(f"LinkedIn: {personal_info.linkedin}")
        if personal_info.website:
            links.append(f"Website: {personal_info.website}")
        
        if links:
            p = doc.add_paragraph(" | ".join(links))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # О себе
        if personal_info.summary:
            doc.add_paragraph()
            doc.add_heading('О себе', level=2)
            doc.add_paragraph(personal_info.summary)
    
    # Образование
    education_list = resume.education.all()
    if education_list:
        doc.add_paragraph()
        doc.add_heading('Образование', level=2)
        for edu in education_list:
            p = doc.add_paragraph()
            p.add_run(f"{edu.institution}\n").bold = True
            p.add_run(f"{edu.degree}, {edu.field_of_study}\n")
            date_str = f"{edu.start_date.strftime('%m.%Y')}"
            if edu.end_date:
                date_str += f" - {edu.end_date.strftime('%m.%Y')}"
            else:
                date_str += " - настоящее время"
            p.add_run(date_str + "\n")
            if edu.description:
                p.add_run(edu.description)
    
    # Опыт работы
    work_experience_list = resume.work_experience.all()
    if work_experience_list:
        doc.add_paragraph()
        doc.add_heading('Опыт работы', level=2)
        for work in work_experience_list:
            p = doc.add_paragraph()
            p.add_run(f"{work.position}\n").bold = True
            p.add_run(f"{work.company}\n")
            date_str = f"{work.start_date.strftime('%m.%Y')}"
            if work.is_current:
                date_str += " - настоящее время"
            elif work.end_date:
                date_str += f" - {work.end_date.strftime('%m.%Y')}"
            p.add_run(date_str + "\n")
            if work.description:
                p.add_run(work.description)
    
    # Навыки
    skills_list = resume.skills.all()
    if skills_list:
        doc.add_paragraph()
        doc.add_heading('Навыки', level=2)
        
        # Группируем навыки по категориям
        skills_by_category = {}
        for skill in skills_list:
            category = skill.get_category_display()
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(f"{skill.name} ({skill.get_level_display()})")
        
        for category, skills in skills_by_category.items():
            p = doc.add_paragraph()
            p.add_run(f"{category}: ").bold = True
            p.add_run(", ".join(skills))
    
    # Достижения
    achievements_list = resume.achievements.all()
    if achievements_list:
        doc.add_paragraph()
        doc.add_heading('Достижения', level=2)
        for achievement in achievements_list:
            p = doc.add_paragraph(style='List Bullet')
            text = achievement.title
            if achievement.date:
                text += f" ({achievement.date.strftime('%Y')})"
            p.add_run(text).bold = True
            if achievement.description:
                p.add_run(f"\n{achievement.description}")
    
    # Языки
    languages_list = resume.languages.all()
    if languages_list:
        doc.add_paragraph()
        doc.add_heading('Языки', level=2)
        lang_strings = [f"{lang.language} - {lang.get_proficiency_level_display()}" 
                       for lang in languages_list]
        doc.add_paragraph(", ".join(lang_strings))
    
    # Сохраняем в BytesIO
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return docx_file